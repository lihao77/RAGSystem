# -*- coding: utf-8 -*-
"""
文档处理工具执行器
实现文档读取、分块、结构化提取等功能
"""
import os
import json
import csv
import difflib
import logging
import re
from typing import Dict, List, Any, Optional
from pathlib import Path

from tools.response_builder import error_result, success_result

logger = logging.getLogger(__name__)


DEFAULT_READ_MAX_LINES = 2000
FILE_SIZE_PREVIEW_THRESHOLD = 32 * 1024  # 32KB

DEFAULT_STRUCTURE_PREVIEW_ROWS = 5
DEFAULT_STRUCTURE_PREVIEW_DEPTH = 3
DEFAULT_STRUCTURE_PREVIEW_FIELDS = 20

_GEOJSON_TYPES = frozenset({
    "FeatureCollection", "Feature",
    "Point", "MultiPoint", "LineString", "MultiLineString",
    "Polygon", "MultiPolygon", "GeometryCollection",
})

_WKT_PATTERN = re.compile(
    r"^\s*(?:SRID\s*=\s*\d+\s*;\s*)?"
    r"(POINT|LINESTRING|POLYGON|MULTIPOINT|MULTILINESTRING|MULTIPOLYGON|GEOMETRYCOLLECTION)"
    r"\s*\(",
    re.IGNORECASE,
)


def _is_geojson(value: Any) -> bool:
    """检测 dict 是否为 GeoJSON 对象。"""
    return isinstance(value, dict) and value.get("type") in _GEOJSON_TYPES


def _count_coords(coords: Any) -> int:
    """递归统计坐标点数量。"""
    if not isinstance(coords, list) or len(coords) == 0:
        return 0
    if isinstance(coords[0], (int, float)):
        return 1
    return sum(_count_coords(c) for c in coords)


def _bbox_from_coords(coords: Any) -> Optional[List[float]]:
    """递归提取坐标的 bbox [min_x, min_y, max_x, max_y]。"""
    if not isinstance(coords, list) or len(coords) == 0:
        return None
    if isinstance(coords[0], (int, float)):
        x, y = float(coords[0]), float(coords[1]) if len(coords) > 1 else 0.0
        return [x, y, x, y]
    bbox = None
    for c in coords:
        sub = _bbox_from_coords(c)
        if sub is None:
            continue
        if bbox is None:
            bbox = list(sub)
        else:
            bbox[0] = min(bbox[0], sub[0])
            bbox[1] = min(bbox[1], sub[1])
            bbox[2] = max(bbox[2], sub[2])
            bbox[3] = max(bbox[3], sub[3])
    return bbox


def _preview_geometry(geom: dict) -> dict:
    """预览单个 GeoJSON 几何体。"""
    coords = geom.get("coordinates", [])
    result: Dict[str, Any] = {"geometry_type": geom.get("type", "unknown")}
    result["coordinate_count"] = _count_coords(coords)
    bbox = geom.get("bbox") or _bbox_from_coords(coords)
    if bbox:
        result["bbox"] = [round(v, 6) for v in bbox]
    return result


def _preview_geojson(value: dict, sample_size: int = 3) -> dict:
    """为 GeoJSON 对象生成紧凑预览，绝不展开 coordinates。"""
    geojson_type = value.get("type", "")

    if geojson_type == "FeatureCollection":
        features = value.get("features", [])
        # 几何类型分布
        type_counts: Dict[str, int] = {}
        total_coords = 0
        merged_bbox: Optional[List[float]] = None
        for f in features:
            geom = f.get("geometry") or {}
            gt = geom.get("type", "null")
            type_counts[gt] = type_counts.get(gt, 0) + 1
            coords = geom.get("coordinates", [])
            total_coords += _count_coords(coords)
            fb = _bbox_from_coords(coords)
            if fb:
                if merged_bbox is None:
                    merged_bbox = list(fb)
                else:
                    merged_bbox[0] = min(merged_bbox[0], fb[0])
                    merged_bbox[1] = min(merged_bbox[1], fb[1])
                    merged_bbox[2] = max(merged_bbox[2], fb[2])
                    merged_bbox[3] = max(merged_bbox[3], fb[3])
        # 属性字段采样
        prop_fields: List[str] = []
        sample_props: List[dict] = []
        for f in features[:sample_size]:
            props = f.get("properties") or {}
            if not prop_fields:
                prop_fields = list(props.keys())[:DEFAULT_STRUCTURE_PREVIEW_FIELDS]
            sampled = {}
            for k in list(props.keys())[:8]:
                v = props[k]
                if isinstance(v, (dict, list)):
                    continue
                # WKT / 长坐标字符串截断
                if isinstance(v, str) and len(v) > 60:
                    if _is_wkt_geometry(v):
                        v = _preview_wkt(v)
                    else:
                        v = _truncate_preview_text(v, limit=80)
                sampled[k] = v
            sample_props.append(sampled)

        result: Dict[str, Any] = {
            "type": "geojson",
            "geojson_type": "FeatureCollection",
            "feature_count": len(features),
            "geometry_types": type_counts,
            "total_coordinates_estimate": total_coords,
            "properties_fields": prop_fields,
        }
        bbox = value.get("bbox") or (merged_bbox if merged_bbox else None)
        if bbox:
            result["bbox"] = [round(v, 6) for v in bbox]
        if sample_props:
            result["sample_properties"] = sample_props
        return result

    if geojson_type == "Feature":
        geom = value.get("geometry") or {}
        props = value.get("properties") or {}
        result = {
            "type": "geojson",
            "geojson_type": "Feature",
        }
        result.update(_preview_geometry(geom))
        result["properties_fields"] = list(props.keys())[:DEFAULT_STRUCTURE_PREVIEW_FIELDS]
        return result

    # 裸几何体 (Point / Polygon / ...)
    result = {"type": "geojson"}
    result.update(_preview_geometry(value))
    return result


def _is_wkt_geometry(value: str) -> bool:
    """检测字符串是否为 WKT 几何体。"""
    return bool(_WKT_PATTERN.match(value))


def _preview_wkt(value: str) -> dict:
    """为 WKT 字符串生成紧凑预览。"""
    m = _WKT_PATTERN.match(value)
    geom_type = m.group(1).capitalize() if m else "Unknown"
    return {
        "type": "wkt_geometry",
        "geometry_type": geom_type,
        "length": len(value),
        "example": value[:60] + ("..." if len(value) > 60 else ""),
    }


def _truncate_preview_text(value: str, limit: int = 120) -> str:
    """Keep text previews compact for agent-friendly structure summaries."""
    if len(value) <= limit:
        return value
    return value[:limit] + "..."


def _format_binary_size(size_bytes: int) -> str:
    """Return a compact binary size label for UI and summaries."""
    if size_bytes < 1024:
        return f"{size_bytes}B"
    if size_bytes < 1024 * 1024:
        value = size_bytes / 1024
        return f"{value:.0f}KB" if value.is_integer() else f"{value:.1f}KB"
    value = size_bytes / (1024 * 1024)
    return f"{value:.1f}MB"


def _scalar_type_name(value: Any) -> str:
    """Return a stable scalar type name."""
    if value is None:
        return "null"
    if isinstance(value, bool):
        return "boolean"
    if isinstance(value, int) and not isinstance(value, bool):
        return "integer"
    if isinstance(value, float):
        return "number"
    if isinstance(value, str):
        return "string"
    return type(value).__name__


def _preview_scalar(value: Any) -> Dict[str, Any]:
    """Build a compact preview for scalar values."""
    if isinstance(value, str) and len(value) > 30 and _is_wkt_geometry(value):
        return _preview_wkt(value)
    preview = {"type": _scalar_type_name(value)}
    if isinstance(value, str):
        preview["example"] = _truncate_preview_text(value)
        preview["length"] = len(value)
    elif value is not None:
        preview["example"] = value
    return preview


def _preview_data_value(
    value: Any,
    *,
    depth: int = 0,
    max_depth: int = DEFAULT_STRUCTURE_PREVIEW_DEPTH,
    max_fields: int = DEFAULT_STRUCTURE_PREVIEW_FIELDS,
    sample_size: int = DEFAULT_STRUCTURE_PREVIEW_ROWS,
) -> Dict[str, Any]:
    """Infer a compact nested structure preview for JSON/YAML-like data."""
    if isinstance(value, dict) and _is_geojson(value):
        return _preview_geojson(value, sample_size=sample_size)

    if depth >= max_depth:
        if isinstance(value, dict):
            return {"type": "object", "key_count": len(value), "truncated": True}
        if isinstance(value, list):
            return {"type": "array", "length": len(value), "truncated": True}
        return _preview_scalar(value)

    if isinstance(value, dict):
        preview_fields = {}
        keys = list(value.keys())
        for key in keys[:max_fields]:
            preview_fields[str(key)] = _preview_data_value(
                value[key],
                depth=depth + 1,
                max_depth=max_depth,
                max_fields=max_fields,
                sample_size=sample_size,
            )

        result = {
            "type": "object",
            "key_count": len(value),
            "keys": [str(key) for key in keys[:max_fields]],
            "fields": preview_fields,
        }
        if len(keys) > max_fields:
            result["truncated_keys"] = len(keys) - max_fields
        return result

    if isinstance(value, list):
        sample_items = value[:sample_size]
        item_types = sorted({_scalar_type_name(item) if not isinstance(item, (dict, list)) else ("object" if isinstance(item, dict) else "array") for item in sample_items})
        result = {
            "type": "array",
            "length": len(value),
            "item_types": item_types,
            "sample_item_count": len(sample_items),
        }

        if not sample_items:
            return result

        if all(isinstance(item, dict) for item in sample_items):
            field_summaries = {}
            field_order = []
            for item in sample_items:
                for key, item_value in item.items():
                    key = str(key)
                    if key not in field_summaries:
                        if len(field_order) >= max_fields:
                            continue
                        field_order.append(key)
                        field_summaries[key] = {
                            "types": set(),
                            "present_in": 0,
                            "example": None,
                        }
                    if key not in field_summaries:
                        continue
                    summary = field_summaries[key]
                    if isinstance(item_value, dict) and _is_geojson(item_value):
                        summary["types"].add("geojson")
                    elif isinstance(item_value, dict):
                        summary["types"].add("object")
                    elif isinstance(item_value, list):
                        summary["types"].add("array")
                    elif isinstance(item_value, str) and len(item_value) > 30 and _is_wkt_geometry(item_value):
                        summary["types"].add("wkt_geometry")
                    else:
                        summary["types"].add(_scalar_type_name(item_value))
                    summary["present_in"] += 1
                    if summary["example"] is None:
                        if isinstance(item_value, dict) and _is_geojson(item_value):
                            summary["example"] = _preview_geojson(item_value)
                        elif isinstance(item_value, (dict, list)):
                            summary["example"] = _preview_data_value(
                                item_value,
                                depth=depth + 1,
                                max_depth=max_depth,
                                max_fields=max_fields,
                                sample_size=sample_size,
                            )
                        else:
                            scalar_preview = _preview_scalar(item_value)
                            summary["example"] = scalar_preview if scalar_preview.get("type") == "wkt_geometry" else scalar_preview.get("example", item_value)

            result["item_structure"] = {
                "type": "object",
                "fields": {
                    key: {
                        "types": sorted(value["types"]),
                        "present_in_sample": value["present_in"],
                        "example": value["example"],
                    }
                    for key, value in field_summaries.items()
                },
            }
            if len(field_order) >= max_fields:
                result["item_structure"]["truncated_fields"] = True
            return result

        result["sample_items"] = [
            _preview_data_value(
                item,
                depth=depth + 1,
                max_depth=max_depth,
                max_fields=max_fields,
                sample_size=sample_size,
            )
            for item in sample_items
        ]
        return result

    return _preview_scalar(value)


def _detect_csv_delimiter(file_path: Path, encoding: str, default: str) -> str:
    """Infer CSV delimiter from a small text sample."""
    try:
        with open(file_path, "r", encoding=encoding, errors="replace", newline="") as handle:
            sample = handle.read(2048)
        if not sample.strip():
            return default
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
        return dialect.delimiter
    except Exception:
        return default


def _infer_csv_column_types(sample_rows: List[Dict[str, Any]], fieldnames: List[str]) -> Dict[str, Any]:
    """Infer coarse types for CSV/TSV columns from sampled rows."""
    column_summary = {}
    for field in fieldnames:
        observed_types = set()
        examples = []
        non_empty_count = 0
        for row in sample_rows:
            raw_value = (row.get(field) or "").strip()
            if raw_value == "":
                continue
            non_empty_count += 1
            if len(examples) < 2:
                examples.append(_truncate_preview_text(raw_value))
            lowered = raw_value.lower()
            if lowered in {"true", "false"}:
                observed_types.add("boolean")
                continue
            try:
                int(raw_value)
                observed_types.add("integer")
                continue
            except ValueError:
                pass
            try:
                float(raw_value)
                observed_types.add("number")
                continue
            except ValueError:
                pass
            observed_types.add("string")

        column_summary[field] = {
            "types": sorted(observed_types) if observed_types else ["string"],
            "non_empty_in_sample": non_empty_count,
            "examples": examples,
        }
    return column_summary


def _preview_delimited_file(
    file_path: Path,
    *,
    encoding: str,
    delimiter: str,
    max_rows: int,
) -> Dict[str, Any]:
    """Preview CSV/TSV structure without materializing the whole file."""
    with open(file_path, "r", encoding=encoding, errors="replace", newline="") as handle:
        reader = csv.DictReader(handle, delimiter=delimiter)
        fieldnames = list(reader.fieldnames or [])
        sample_rows = []
        total_rows = 0
        for row in reader:
            total_rows += 1
            if len(sample_rows) < max_rows:
                sample_rows.append({key: row.get(key, "") for key in fieldnames})

    return {
        "root_type": "table",
        "delimiter": delimiter,
        "column_count": len(fieldnames),
        "columns": fieldnames,
        "sample_row_count": len(sample_rows),
        "total_rows": total_rows,
        "column_types": _infer_csv_column_types(sample_rows, fieldnames),
        "sample_rows": sample_rows,
    }


def _preview_text_file(file_path: Path, *, encoding: str, max_rows: int) -> Dict[str, Any]:
    """Preview basic structure for plain text-like files."""
    with open(file_path, "r", encoding=encoding, errors="replace") as handle:
        lines = handle.readlines()

    preview_lines = [line.rstrip("\n\r") for line in lines[:max_rows]]
    non_empty_lines = [line for line in lines if line.strip()]
    max_line_length = max((len(line.rstrip("\n\r")) for line in lines), default=0)
    avg_line_length = (
        round(sum(len(line.rstrip("\n\r")) for line in lines) / len(lines), 2)
        if lines else 0
    )

    return {
        "root_type": "text",
        "total_lines": len(lines),
        "non_empty_lines": len(non_empty_lines),
        "max_line_length": max_line_length,
        "average_line_length": avg_line_length,
        "preview_lines": preview_lines,
    }


def _load_structured_document(file_path: Path, *, encoding: str) -> Any:
    """Load JSON/YAML-like data files."""
    suffix = file_path.suffix.lower()
    if suffix == ".json":
        with open(file_path, "r", encoding=encoding, errors="replace") as handle:
            text = handle.read()
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            import ast
            return ast.literal_eval(text)
    if suffix in {".yaml", ".yml"}:
        try:
            import yaml
        except ImportError as error:
            raise RuntimeError("需要安装 PyYAML 才能预览 YAML 结构") from error
        with open(file_path, "r", encoding=encoding, errors="replace") as handle:
            return yaml.safe_load(handle)
    raise ValueError(f"不支持的结构化文档格式: {suffix}")


def read_document(file_path: str, encoding: str = "utf-8"):
    """读取文档文件内容"""
    file_path = Path(file_path)

    if not file_path.exists():
        return error_result(f"文件不存在: {file_path}", tool_name="read_document")

    suffix = file_path.suffix.lower()

    try:
        # TXT/MD 文件
        if suffix in ['.txt', '.md']:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            return success_result(
                content=content,
                summary=f"文档读取成功: {file_path.name}",
                output_type="text",
                metadata={
                    "file_type": suffix[1:],
                    "char_count": len(content),
                    "file_path": str(file_path),
                },
                tool_name="read_document",
            )

        # PDF 文件
        elif suffix == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                return success_result(
                    content=content,
                    summary=f"PDF 读取成功: {file_path.name}",
                    output_type="text",
                    metadata={
                        "file_type": "pdf",
                        "page_count": len(reader.pages),
                        "char_count": len(content),
                        "file_path": str(file_path),
                    },
                    tool_name="read_document",
                )
            except ImportError:
                return error_result("需要安装 PyPDF2: pip install PyPDF2", tool_name="read_document")

        # Word 文件
        elif suffix in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                return success_result(
                    content=content,
                    summary=f"Word 文档读取成功: {file_path.name}",
                    output_type="text",
                    metadata={
                        "file_type": "docx",
                        "paragraph_count": len(doc.paragraphs),
                        "char_count": len(content),
                        "file_path": str(file_path),
                    },
                    tool_name="read_document",
                )
            except ImportError:
                return error_result("需要安装 python-docx: pip install python-docx", tool_name="read_document")

        else:
            return error_result(f"不支持的文件格式: {suffix}", tool_name="read_document")

    except Exception as e:
        return error_result(f"读取文件失败: {str(e)}", tool_name="read_document")


def chunk_document(
    content: str,
    chunk_size: int = 2000,
    chunk_overlap: int = 200,
    strategy: str = "fixed"
):
    """将长文档分块"""
    try:
        chunks = []

        if strategy == "fixed":
            # 固定大小分块
            start = 0
            chunk_id = 0
            while start < len(content):
                end = start + chunk_size
                chunk_text = content[start:end]
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": chunk_text,
                    "start_pos": start,
                    "end_pos": min(end, len(content)),
                    "char_count": len(chunk_text)
                })
                start = end - chunk_overlap
                chunk_id += 1

        elif strategy == "paragraph":
            # 按段落分块
            paragraphs = content.split('\n\n')
            current_chunk = ""
            chunk_id = 0
            start_pos = 0

            for para in paragraphs:
                if len(current_chunk) + len(para) > chunk_size and current_chunk:
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": current_chunk.strip(),
                        "start_pos": start_pos,
                        "end_pos": start_pos + len(current_chunk),
                        "char_count": len(current_chunk)
                    })
                    chunk_id += 1
                    start_pos += len(current_chunk)
                    current_chunk = para + "\n\n"
                else:
                    current_chunk += para + "\n\n"

            # 最后一块
            if current_chunk:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk),
                    "char_count": len(current_chunk)
                })

        elif strategy == "semantic":
            # 语义分块（简化版：按句子分块）
            import re
            sentences = re.split(r'([。！？\n])', content)
            current_chunk = ""
            chunk_id = 0
            start_pos = 0

            for i in range(0, len(sentences), 2):
                sentence = sentences[i] + (sentences[i+1] if i+1 < len(sentences) else "")
                if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": current_chunk.strip(),
                        "start_pos": start_pos,
                        "end_pos": start_pos + len(current_chunk),
                        "char_count": len(current_chunk)
                    })
                    chunk_id += 1
                    start_pos += len(current_chunk)
                    current_chunk = sentence
                else:
                    current_chunk += sentence

            if current_chunk:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk),
                    "char_count": len(current_chunk)
                })

        return success_result(
            content=chunks,
            summary=f"文档分块成功，共 {len(chunks)} 块",
            output_type="json",
            metadata={
                "total_chunks": len(chunks),
                "strategy": strategy,
                "chunk_size": chunk_size,
                "chunk_overlap": chunk_overlap,
            },
            tool_name="chunk_document",
        )

    except Exception as e:
        return error_result(f"分块失败: {str(e)}", tool_name="chunk_document")


def extract_structured_data(
    text: str,
    schema: Dict[str, Any],
    instruction: Optional[str] = None,
    examples: Optional[List[Dict]] = None
):
    """从文本中提取结构化数据（使用 LLM）"""
    try:
        from model_adapter import get_default_adapter

        adapter = get_default_adapter()

        # 构建提示词
        prompt = f"""请从以下文本中提取结构化信息，严格按照 JSON Schema 格式返回。

JSON Schema:
{json.dumps(schema, ensure_ascii=False, indent=2)}
"""

        if instruction:
            prompt += f"\n提取要求：\n{instruction}\n"

        if examples:
            prompt += f"\n示例格式：\n{json.dumps(examples[0], ensure_ascii=False, indent=2)}\n"

        prompt += f"\n文本内容：\n{text}\n\n请直接返回 JSON 格式的提取结果，不要包含任何其他说明文字。"

        # 调用 LLM（使用 JSON mode）
        response = adapter.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )

        # 解析结果
        result_text = response.get("content", "")
        extracted_data = json.loads(result_text)

        return success_result(
            content=extracted_data,
            summary="结构化提取成功",
            output_type="json",
            metadata={"text_length": len(text)},
            tool_name="extract_structured_data",
        )

    except Exception as e:
        return error_result(f"提取失败: {str(e)}", tool_name="extract_structured_data")


def merge_extracted_data(
    data_list: List[Dict],
    merge_strategy: str = "deduplicate",
    unique_key: Optional[str] = None
):
    """合并多个提取结果"""
    try:
        if merge_strategy == "append":
            # 简单追加
            merged = []
            for data in data_list:
                if isinstance(data, list):
                    merged.extend(data)
                else:
                    merged.append(data)

        elif merge_strategy == "deduplicate":
            # 去重
            if not unique_key:
                return error_result("去重策略需要指定 unique_key", tool_name="merge_extracted_data")

            seen = set()
            merged = []
            for data in data_list:
                items = data if isinstance(data, list) else [data]
                for item in items:
                    key_value = item.get(unique_key)
                    if key_value and key_value not in seen:
                        seen.add(key_value)
                        merged.append(item)

        elif merge_strategy == "merge_by_key":
            # 按键合并（相同键的数据合并）
            if not unique_key:
                return error_result("按键合并策略需要指定 unique_key", tool_name="merge_extracted_data")

            merged_dict = {}
            for data in data_list:
                items = data if isinstance(data, list) else [data]
                for item in items:
                    key_value = item.get(unique_key)
                    if key_value:
                        if key_value in merged_dict:
                            # 合并字段
                            merged_dict[key_value].update(item)
                        else:
                            merged_dict[key_value] = item
            merged = list(merged_dict.values())

        return success_result(
            content=merged,
            summary=f"提取结果合并成功，共 {len(merged)} 项",
            output_type="json",
            metadata={
                "total_items": len(merged),
                "merge_strategy": merge_strategy,
                "unique_key": unique_key,
            },
            tool_name="merge_extracted_data",
        )

    except Exception as e:
        return error_result(f"合并失败: {str(e)}", tool_name="merge_extracted_data")


def preview_data_structure(
    file_path: str,
    encoding: str = "utf-8",
    max_preview_rows: int = DEFAULT_STRUCTURE_PREVIEW_ROWS,
    max_depth: int = DEFAULT_STRUCTURE_PREVIEW_DEPTH,
    max_fields: int = DEFAULT_STRUCTURE_PREVIEW_FIELDS,
) -> Any:
    """预览常见数据文件的数据结构，帮助 Agent 决定下一步读取或提取策略。"""
    try:
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            return error_result(f"文件不存在: {file_path}", tool_name="preview_data_structure")

        if max_preview_rows < 1:
            return error_result("max_preview_rows 必须 >= 1", tool_name="preview_data_structure")
        if max_depth < 1:
            return error_result("max_depth 必须 >= 1", tool_name="preview_data_structure")
        if max_fields < 1:
            return error_result("max_fields 必须 >= 1", tool_name="preview_data_structure")

        suffix = file_path_obj.suffix.lower()
        file_size = file_path_obj.stat().st_size

        if suffix in {".json", ".yaml", ".yml"}:
            data = _load_structured_document(file_path_obj, encoding=encoding)
            structure = _preview_data_value(
                data,
                max_depth=max_depth,
                max_fields=max_fields,
                sample_size=max_preview_rows,
            )
            file_type = suffix[1:]
        elif suffix in {".csv", ".tsv"}:
            delimiter = "\t" if suffix == ".tsv" else _detect_csv_delimiter(file_path_obj, encoding, ",")
            structure = _preview_delimited_file(
                file_path_obj,
                encoding=encoding,
                delimiter=delimiter,
                max_rows=max_preview_rows,
            )
            file_type = "tsv" if delimiter == "\t" else "csv"
        else:
            structure = _preview_text_file(
                file_path_obj,
                encoding=encoding,
                max_rows=max_preview_rows,
            )
            file_type = suffix[1:] if suffix else "text"

        content = {
            "file_path": str(file_path_obj),
            "file_name": file_path_obj.name,
            "file_type": file_type,
            "file_size": file_size,
            "structure": structure,
        }
        return success_result(
            content=content,
            summary=f"已预览文件数据结构: {file_path_obj.name}",
            output_type="json",
            metadata={
                "file_path": str(file_path_obj),
                "file_type": file_type,
                "file_size": file_size,
                "max_preview_rows": max_preview_rows,
                "max_depth": max_depth,
                "max_fields": max_fields,
            },
            tool_name="preview_data_structure",
        )
    except Exception as e:
        return error_result(f"预览数据结构失败: {str(e)}", tool_name="preview_data_structure")


def write_file(
    content: str,
    file_path: Optional[str] = None,
    encoding: str = "utf-8",
    mode: str = "text",
    *,
    session_id: Optional[str] = None,
    run_id: Optional[str] = None,
    default_output_space: Optional[str] = None,
    workspace_root: Optional[str] = None,
) -> Any:
    """写入文本内容到文件。JSON 请先用 json.dumps 序列化为字符串再传入。

    file_path 由 dispatcher 预处理层保证为绝对路径（未指定时已预分配）。
    """
    try:
        if not file_path:
            from tools.path_resolution import resolve_managed_path
            suffix = '.json' if mode == 'json' else '.txt'
            file_path = str(resolve_managed_path(
                None,
                session_id=session_id,
                run_id=run_id,
                caller='direct',
                operation='write',
                default_output_space=default_output_space,
                workspace_root=workspace_root,
                suffix=suffix,
            ))
        file_path_obj = Path(file_path)
        if not file_path_obj.is_absolute():
            logger.warning("write_file 收到非绝对路径，尝试 resolve: %s", file_path)
            file_path_obj = file_path_obj.resolve()

        dir_path = file_path_obj.parent
        if dir_path:
            dir_path.mkdir(parents=True, exist_ok=True)

        with open(file_path_obj, 'w', encoding=encoding) as f:
            if mode == "json":
                if isinstance(content, str):
                    try:
                        json.dump(json.loads(content), f, ensure_ascii=False, indent=2)
                    except json.JSONDecodeError:
                        f.write(content)
                else:
                    json.dump(content, f, ensure_ascii=False, indent=2)
            else:
                f.write(content if isinstance(content, str) else str(content))

        file_size = file_path_obj.stat().st_size
        from tools.path_resolution import to_display_path
        display = to_display_path(file_path_obj)
        return success_result(
            content={"file_path": str(file_path_obj), "file_size": file_size, "display_path": display},
            summary=f"文件已写入: {display}（{file_size} 字节）",
            output_type="text",
            tool_name="write_file",
        )

    except Exception as e:
        return error_result(f"写入文件失败: {str(e)}", tool_name="write_file")


def read_file(
    file_path: str,
    encoding: str = "utf-8",
    offset: int = 1,
    limit: int = DEFAULT_READ_MAX_LINES,
    *,
    caller: str = "direct",
    event_bus=None,
    session_id: Optional[str] = None,
    workspace_root: Optional[str] = None,
    run_id: Optional[str] = None,
    _skip_preview: bool = False,
) -> Any:
    """按行号读取文件内容，返回原始文本内容。支持大文件预览确认（仅 direct 调用）。

    file_path 由 dispatcher 预处理层保证为绝对路径。
    """
    try:
        from tools.path_resolution import resolve_managed_path
        resolved_path = resolve_managed_path(
            file_path,
            session_id=session_id,
            run_id=run_id,
            caller=caller,
            operation='read',
            workspace_root=workspace_root,
        )
        file_path_obj = Path(resolved_path)

        if not file_path_obj.exists():
            return error_result(f"文件不存在: {file_path}", tool_name="read_file")

        if offset < 1:
            return error_result("offset 必须 >= 1", tool_name="read_file")
        if limit < 1:
            return error_result("limit 必须 >= 1", tool_name="read_file")

        file_size = file_path_obj.stat().st_size
        _user_approved_full_read = False

        # Agent 显式指定了读取范围 → 按需读取，跳过大文件确认
        agent_specified_range = (offset != 1 or limit != DEFAULT_READ_MAX_LINES)

        # 大文件预览确认：仅非代码执行调用、非 _skip_preview、有 event_bus/session_id 且非按需读取时触发
        if (
            not _skip_preview
            and caller != "code_execution"
            and file_size > FILE_SIZE_PREVIEW_THRESHOLD
            and not agent_specified_range
            and event_bus is not None
            and session_id is not None
        ):
            # 读取前 N KB 作为预览
            with open(file_path_obj, "r", encoding=encoding, errors="replace") as f:
                preview_raw = f.read(FILE_SIZE_PREVIEW_THRESHOLD)

            preview_lines = preview_raw.splitlines(keepends=True)
            preview_formatted = preview_raw.rstrip("\n")

            # 复用审批机制发布确认事件
            import uuid as _uuid
            from agents.events import Event, EventType
            from agents.task_registry import get_task_registry

            approval_id = str(_uuid.uuid4())
            registry = get_task_registry()
            wait_evt = registry.add_pending_approval(session_id, approval_id)

            event_bus.publish(Event(
                type=EventType.USER_APPROVAL_REQUIRED,
                session_id=session_id,
                data={
                    "approval_id": approval_id,
                    "approval_type": "file_read_confirm",
                    "tool_name": "read_file",
                    "file_path": str(file_path_obj),
                    "file_size": file_size,
                    "preview_threshold": FILE_SIZE_PREVIEW_THRESHOLD,
                    "preview": preview_formatted,
                    "description": f"文件 {file_path_obj.name} 较大（{file_size} 字节），是否读取完整内容？",
                }
            ))

            if wait_evt is not None:
                from utils.timeout_pause import pause_current, resume_current

                pause_current()
                try:
                    wait_evt.wait()
                finally:
                    resume_current()
                approved, _ = registry.get_approval_result(session_id, approval_id)
                if not approved:
                    # 用户拒绝 → 返回预览内容
                    return success_result(
                        content=preview_formatted,
                        summary=(
                            f"文件预览: {file_path}"
                            f"（仅前 {_format_binary_size(FILE_SIZE_PREVIEW_THRESHOLD)}，{file_size} 字节总计）"
                        ),
                        output_type="text",
                        metadata={
                            "file_path": str(file_path_obj),
                            "file_size": file_size,
                            "preview_only": True,
                            "preview_threshold": FILE_SIZE_PREVIEW_THRESHOLD,
                            "preview_lines": len(preview_lines),
                        },
                        tool_name="read_file",
                    )
                # 用户批准 → 标记，后续 metadata 会携带此标记
                _user_approved_full_read = True

        # 正式读取
        with open(file_path_obj, "r", encoding=encoding, errors="replace") as f:
            all_lines = f.readlines()

        total_lines = len(all_lines)
        start_idx = offset - 1  # 0-based index
        end_idx = min(start_idx + limit, total_lines)

        if start_idx >= total_lines:
            return success_result(
                content="",
                summary=f"offset {offset} 超出文件总行数 {total_lines}",
                output_type="text",
                metadata={
                    "file_path": str(file_path_obj),
                    "file_size": file_size,
                    "total_lines": total_lines,
                    "start_line": offset,
                    "end_line": offset,
                    "has_more": False,
                    "next_offset": None,
                },
                tool_name="read_file",
            )

        selected_lines = all_lines[start_idx:end_idx]
        # 返回原始内容，不再对单行长度做截断
        content = "".join(selected_lines)
        content = content.rstrip("\n")

        has_more = end_idx < total_lines
        actual_end_line = start_idx + len(selected_lines)
        next_offset = actual_end_line + 1 if has_more else None

        summary = (
            f"文件读取成功: {file_path}（行 {offset}-{actual_end_line}，"
            f"共 {total_lines} 行，{file_size} 字节）"
        )
        if has_more:
            summary += f"；还有后续内容，可继续调用 read_file(offset={next_offset})"
        else:
            summary += "；已到文件末尾"

        return success_result(
            content=content,
            summary=summary,
            output_type="text",
            metadata={
                "file_path": str(file_path_obj),
                "file_size": file_size,
                "total_lines": total_lines,
                "start_line": offset,
                "end_line": actual_end_line,
                "has_more": has_more,
                "next_offset": next_offset,
                "user_approved_full_read": _user_approved_full_read,
            },
            tool_name="read_file",
        )

    except Exception as e:
        return error_result(f"读取文件失败: {str(e)}", tool_name="read_file")


def edit_file(
    file_path: str,
    old_string: str,
    new_string: str,
    encoding: str = "utf-8",
    replace_all: bool = False,
    *,
    workspace_root: Optional[str] = None,
    session_id: Optional[str] = None,
    run_id: Optional[str] = None,
    caller: str = 'direct',
) -> Any:
    """精准字符串替换编辑文件。old_string 必须唯一匹配（除非 replace_all=True）。

    file_path 由 dispatcher 预处理层保证为绝对路径。
    """
    try:
        from tools.path_resolution import resolve_managed_path
        resolved_path = resolve_managed_path(
            file_path,
            session_id=session_id,
            run_id=run_id,
            caller=caller,
            operation='edit',
            workspace_root=workspace_root,
        )
        file_path_obj = Path(resolved_path)

        if not file_path_obj.exists():
            return error_result(f"文件不存在: {file_path}", tool_name="edit_file")

        with open(file_path_obj, "r", encoding=encoding) as f:
            content = f.read()

        count = content.count(old_string)

        if count == 0:
            return error_result(
                f"未找到匹配内容，请检查 old_string 是否与文件内容完全一致",
                tool_name="edit_file",
            )

        if count > 1 and not replace_all:
            return error_result(
                f"匹配不唯一（找到 {count} 处匹配）。请提供更多上下文使 old_string 唯一，或设 replace_all=true",
                tool_name="edit_file",
            )

        # 执行替换
        if replace_all:
            new_content = content.replace(old_string, new_string)
            replacements = count
        else:
            new_content = content.replace(old_string, new_string, 1)
            replacements = 1

        # 写回文件
        with open(file_path_obj, "w", encoding=encoding) as f:
            f.write(new_content)

        # 构建 unified diff 预览
        old_lines = content.splitlines(keepends=True)
        new_lines = new_content.splitlines(keepends=True)
        diff = difflib.unified_diff(
            old_lines, new_lines,
            fromfile=f"a/{file_path_obj.name}",
            tofile=f"b/{file_path_obj.name}",
        )
        diff_preview = "".join(diff)
        if len(diff_preview) > 2000:
            diff_preview = diff_preview[:2000] + "\n... [DIFF TRUNCATED]"

        new_size = file_path_obj.stat().st_size

        return success_result(
            content={
                "file_path": str(file_path_obj),
                "replacements": replacements,
                "file_size": new_size,
                "diff_preview": diff_preview,
            },
            summary=f"文件编辑成功: {file_path}（替换 {replacements} 处，{new_size} 字节）",
            output_type="json",
            metadata={
                "file_path": str(file_path_obj),
                "replacements": replacements,
                "file_size": new_size,
            },
            tool_name="edit_file",
        )

    except Exception as e:
        return error_result(f"编辑文件失败: {str(e)}", tool_name="edit_file")


